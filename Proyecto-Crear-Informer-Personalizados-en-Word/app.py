import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import io

# Función para crear el reporte del informe
def crear_reporte(template_path, data, chart_data=None):
    st.write("Iniciando la creación del informe...")
    doc = Document(template_path)

    # Recorremos los párrafos del documento
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f'{{{{{key}}}}}' in paragraph.text:
                st.write(f"Reemplazando {key} con {value} en el informe.")
            paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))

    # Lógica para el gráfico
    if chart_data is not None:
        st.write("Generando gráfico...")
        
        # Crear gráfico con matplotlib => plt
        plt.figure(figsize=(6, 4))
        plt.bar(chart_data['labels'], chart_data['values'])
        plt.title(chart_data['title'])
        plt.xlabel(chart_data.get('xlabel', ''))
        plt.ylabel(chart_data.get('ylabel', ''))
        
        # Guardadamos el gráfico en un buffer
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png')
        img_buffer.seek(0)
        
        st.write("Insertando gráfico en el marcador del documento...")
        
        # Insertamos gráfico en el documento
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if '[Aqui se insertará el gráfico]' in run.text:
                    run.text = run.text.replace('[Aqui se insertará el gráfico]', '')
                    run.add_picture(img_buffer, width=Inches(6))

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    st.write("Informe creado con éxito.")
    return output

# Función Principal
def main():
    st.header("Generador de Informe Automatizado")

    template_file = st.file_uploader("Cargar Plantilla Word", type="docx")
    data_file = st.file_uploader("Cargar Datos", type=["xlsx", "csv"])
    
    if template_file and data_file:
        st.success("Los archivos se cargaron correctamente.")
        data_frame = pd.read_csv(data_file) if data_file.name.endswith('csv') else pd.read_excel(data_file)

        st.subheader("Datos Cargados")
        st.dataframe(data_frame)

        index = st.selectbox("Seleccionar fila para el informe", options=range(len(data_frame)))
        selected_data = data_frame.iloc[index].to_dict()
        
        # Opcional: gráfico
        generar_chart = st.checkbox("¿Deseas generar un gráfico?")
        chart_data = None
        
        if generar_chart:
            chart_title = st.text_input("Título del gráfico", "Gráfico de Datos")
            column_x = st.selectbox("Columna para el eje X", options=data_frame.columns)
            column_y = st.selectbox("Columna para el eje Y", options=data_frame.columns)

            chart_data = {
                'title': chart_title,
                'labels': data_frame[column_x].tolist(), 
                'values': data_frame[column_y].tolist(),
                'xlabel': column_x,
                'ylabel': column_y
            }

            st.write("Datos del gráfico", chart_data)

        if st.button("Generar Informe"):
            output = crear_reporte(template_file, selected_data, chart_data)
            st.download_button("Descargar informe", output, "informe_generado.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

main()
